import pandas as pd
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import os
from datetime import datetime
from pathlib import Path

class MobilePDF(FPDF):
    """Classe PDF personnalisée avec en-têtes et pieds de page"""
    
    def __init__(self, title="Rapport"):
        super().__init__()
        self.report_title = title
        self.set_auto_page_break(auto=True, margin=15)
        
    def header(self):
        """En-tête de page"""
        self.set_font('Arial', 'B', 16)
        self.cell(0, 10, self.report_title, 0, 1, 'C')
        self.set_font('Arial', '', 10)
        self.cell(0, 5, f'Généré le: {datetime.now().strftime("%d/%m/%Y à %H:%M")}', 0, 1, 'C')
        self.ln(10)
        
    def footer(self):
        """Pied de page"""
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')
        
    def add_table(self, headers, data, col_widths=None):
        """
        Ajouter un tableau au PDF
        
        Args:
            headers (list): En-têtes des colonnes
            data (list): Données du tableau
            col_widths (list): Largeurs des colonnes (optionnel)
        """
        if not headers or not data:
            self.set_font('Arial', '', 12)
            self.cell(0, 10, 'Aucune donnée à afficher', 0, 1, 'C')
            return
            
        # Calculer les largeurs de colonnes
        if not col_widths:
            available_width = self.w - 2 * self.l_margin
            col_widths = [available_width / len(headers)] * len(headers)
        
        # En-têtes
        self.set_font('Arial', 'B', 10)
        self.set_fill_color(54, 96, 146)  # Bleu foncé
        self.set_text_color(255, 255, 255)  # Blanc
        
        for i, header in enumerate(headers):
            self.cell(col_widths[i], 8, str(header), 1, 0, 'C', True)
        self.ln()
        
        # Données
        self.set_font('Arial', '', 9)
        self.set_text_color(0, 0, 0)  # Noir
        
        fill = False
        for row in data:
            if fill:
                self.set_fill_color(240, 240, 240)  # Gris clair
            else:
                self.set_fill_color(255, 255, 255)  # Blanc
                
            # Vérifier si on a besoin d'une nouvelle page
            if self.get_y() > self.h - 30:  # 30mm du bas
                self.add_page()
                # Réafficher les en-têtes sur la nouvelle page
                self.set_font('Arial', 'B', 10)
                self.set_fill_color(54, 96, 146)
                self.set_text_color(255, 255, 255)
                for i, header in enumerate(headers):
                    self.cell(col_widths[i], 8, str(header), 1, 0, 'C', True)
                self.ln()
                self.set_font('Arial', '', 9)
                self.set_text_color(0, 0, 0)
            
            for i, cell_data in enumerate(row):
                # Tronquer le texte si trop long
                text = str(cell_data)[:30] + "..." if len(str(cell_data)) > 30 else str(cell_data)
                self.cell(col_widths[i], 6, text, 1, 0, 'C', fill)
            self.ln()
            fill = not fill  # Alterner les couleurs
            
    def add_text(self, text, font_size=12, bold=False):
        """Ajouter du texte simple"""
        font_style = 'B' if bold else ''
        self.set_font('Arial', font_style, font_size)
        self.set_text_color(0, 0, 0)
        self.cell(0, 8, text, 0, 1)
        self.ln(3)

class DataExporter:
    """
    Classe pour exporter des données en format PDF, Word (.docx) et Excel (.xlsx)
    Version optimisée pour Android avec fpdf au lieu de ReportLab
    
    Usage:
        exporter = DataExporter()
        exporter.export_data(data, "mon_fichier", formats=['pdf', 'xlsx', 'docx'])
    """
    
    def __init__(self, output_dir="exports"):
        """
        Initialiser l'exporteur
        
        Args:
            output_dir (str): Répertoire de sortie pour les fichiers
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
    def export_data(self, data, filename_base, formats=['pdf', 'xlsx', 'docx'], 
                   title="Rapport de données", headers=None):
        """
        Exporter des données dans les formats spécifiés
        
        Args:
            data (list): Liste de listes ou liste de dictionnaires
            filename_base (str): Nom de base pour les fichiers (sans extension)
            formats (list): Liste des formats ('pdf', 'xlsx', 'docx')
            title (str): Titre du document
            headers (list): En-têtes des colonnes (optionnel)
        
        Returns:
            dict: Dictionnaire des chemins des fichiers créés
        """
        # Normaliser les données
        normalized_data = self._normalize_data(data, headers)
        
        created_files = {}
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        for format_type in formats:
            filename = f"{filename_base}_{timestamp}.{format_type}"
            filepath = self.output_dir / filename
            
            try:
                if format_type.lower() == 'pdf':
                    self._export_to_pdf(normalized_data, filepath, title)
                elif format_type.lower() in ['xlsx', 'excel']:
                    self._export_to_excel(normalized_data, filepath, title)
                elif format_type.lower() in ['docx', 'word', 'docs']:  # Ajout de 'docs'
                    self._export_to_word(normalized_data, filepath, title)
                else:
                    print(f"Format {format_type} non supporté")
                    continue
                
                created_files[format_type] = str(filepath)
                print(f"✓ Fichier {format_type.upper()} créé: {filepath}")
                
            except Exception as e:
                print(f"✗ Erreur lors de la création du fichier {format_type}: {e}")
                # Afficher plus de détails sur l'erreur pour debug
                import traceback
                traceback.print_exc()
        
        return created_files
    
    def _normalize_data(self, data, headers=None):
        """
        Normaliser les données en format standard
        
        Args:
            data: Données d'entrée (diverses formes)
            headers: En-têtes optionnels
        
        Returns:
            dict: {'headers': [...], 'rows': [[...], [...]]}
        """
        if not data:
            return {'headers': headers or [], 'rows': []}
        
        # Si c'est une liste de dictionnaires
        if isinstance(data[0], dict):
            headers = headers or list(data[0].keys())
            rows = [[str(row.get(col, '')) for col in headers] for row in data]
        
        # Si c'est une liste de listes
        elif isinstance(data[0], (list, tuple)):
            headers = headers or [f"Colonne {i+1}" for i in range(len(data[0]))]
            rows = [[str(cell) for cell in row] for row in data]
        
        # Si c'est une liste simple
        else:
            headers = headers or ['Valeur']
            rows = [[str(item)] for item in data]
        
        return {'headers': headers, 'rows': rows}
    
    def _export_to_pdf(self, normalized_data, filepath, title):
        """Exporter en PDF avec fpdf (compatible Android)"""
        pdf = MobilePDF(title=title)
        pdf.add_page()
        
        # Ajouter des informations
        pdf.add_text(f"Nombre d'enregistrements: {len(normalized_data['rows'])}", bold=True)
        pdf.ln(5)
        
        # Ajouter le tableau
        if normalized_data['rows']:
            # Calculer les largeurs optimales
            headers = normalized_data['headers']
            available_width = pdf.w - 2 * pdf.l_margin
            
            if len(headers) > 0:
                # Largeurs égales par défaut
                col_widths = [available_width / len(headers)] * len(headers)
                
                # Ajuster si certaines colonnes sont plus longues
                max_lengths = []
                for i, header in enumerate(headers):
                    max_len = len(str(header))
                    for row in normalized_data['rows'][:5]:  # Échantillon des 5 premières lignes
                        if i < len(row):
                            max_len = max(max_len, len(str(row[i])[:30]))  # Limité à 30 char
                    max_lengths.append(max_len)
                
                # Redistribuer les largeurs selon le contenu
                total_chars = sum(max_lengths) if sum(max_lengths) > 0 else len(headers)
                col_widths = [(max_len / total_chars) * available_width for max_len in max_lengths]
                
                pdf.add_table(headers, normalized_data['rows'], col_widths)
        else:
            pdf.add_text("Aucune donnée à afficher", font_size=14)
        
        # Sauvegarder
        pdf.output(str(filepath))
    
    def _export_to_excel(self, normalized_data, filepath, title):
        """Exporter en Excel avec pandas et openpyxl"""
        # Créer le DataFrame
        if normalized_data['rows']:
            df = pd.DataFrame(normalized_data['rows'], columns=normalized_data['headers'])
        else:
            df = pd.DataFrame()
        
        # Écrire dans Excel avec mise en forme
        with pd.ExcelWriter(str(filepath), engine='openpyxl') as writer:
            # Feuille principale
            df.to_excel(writer, sheet_name='Données', index=False)
            
            # Accéder au workbook et worksheet pour la mise en forme
            workbook = writer.book
            worksheet = writer.sheets['Données']
            
            # Mise en forme des en-têtes
            try:
                from openpyxl.styles import Font, PatternFill, Alignment
                
                header_font = Font(bold=True, color='FFFFFF')
                header_fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
                
                for col in range(1, len(normalized_data['headers']) + 1):
                    cell = worksheet.cell(row=1, column=col)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = Alignment(horizontal='center')
                
                # Ajuster la largeur des colonnes
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    worksheet.column_dimensions[column_letter].width = adjusted_width
            
            except ImportError:
                print("openpyxl.styles non disponible - export Excel basique")
            
            # Ajouter une feuille d'informations
            info_sheet = workbook.create_sheet('Informations')
            info_sheet['A1'] = 'Titre:'
            info_sheet['B1'] = title
            info_sheet['A2'] = 'Date de création:'
            info_sheet['B2'] = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
            info_sheet['A3'] = 'Nombre de lignes:'
            info_sheet['B3'] = len(normalized_data['rows'])
    
    def _export_to_word(self, normalized_data, filepath, title):
        """Exporter en Word avec python-docx"""
        doc = Document()
        
        # Titre principal
        title_paragraph = doc.add_heading(title, level=1)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Informations
        info_para = doc.add_paragraph()
        info_para.add_run(f"Date de génération: ").bold = True
        info_para.add_run(datetime.now().strftime('%d/%m/%Y à %H:%M'))
        
        info_para2 = doc.add_paragraph()
        info_para2.add_run(f"Nombre d'enregistrements: ").bold = True
        info_para2.add_run(str(len(normalized_data['rows'])))
        
        doc.add_paragraph()  # Espace
        
        # Tableau
        if normalized_data['rows']:
            # Créer le tableau
            table = doc.add_table(
                rows=len(normalized_data['rows']) + 1, 
                cols=len(normalized_data['headers'])
            )
            table.style = 'Table Grid'
            
            # En-têtes
            header_cells = table.rows[0].cells
            for i, header in enumerate(normalized_data['headers']):
                header_cells[i].text = str(header)
                # Mettre en gras les en-têtes
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Données
            for row_idx, row_data in enumerate(normalized_data['rows']):
                row_cells = table.rows[row_idx + 1].cells
                for col_idx, cell_data in enumerate(row_data):
                    if col_idx < len(row_cells):  # Vérification de sécurité
                        row_cells[col_idx].text = str(cell_data)
            
            # Ajustement automatique de la largeur des colonnes
            for column in table.columns:
                for cell in column.cells:
                    cell.width = Inches(2.0)  # Largeur par défaut
        
        else:
            doc.add_paragraph("Aucune donnée à afficher.")
        
        doc.save(str(filepath))

# Fonctions utilitaires simples pour usage rapide
def export_to_pdf(data, filename, title="Rapport", headers=None):
    """Fonction simple pour exporter seulement en PDF"""
    exporter = DataExporter()
    return exporter.export_data(data, filename, formats=['pdf'], title=title, headers=headers)

def export_to_excel(data, filename, title="Rapport", headers=None):
    """Fonction simple pour exporter seulement en Excel"""
    exporter = DataExporter()
    return exporter.export_data(data, filename, formats=['xlsx'], title=title, headers=headers)

def export_to_word(data, filename, title="Rapport", headers=None):
    """Fonction simple pour exporter seulement en Word"""
    exporter = DataExporter()
    return exporter.export_data(data, filename, formats=['docx'], title=title, headers=headers)

# Exemple d'utilisation
if __name__ == "__main__":
    # Données d'exemple sous différentes formes
    
    # 1. Liste de dictionnaires
    data_dict = [
        {'Nom': 'Jean Dupont', 'Age': 30, 'Ville': 'Paris'},
        {'Nom': 'Marie Martin', 'Age': 25, 'Ville': 'Lyon'},
        {'Nom': 'Pierre Durand', 'Age': 35, 'Ville': 'Marseille'}
    ]
    
    # 2. Liste de listes
    data_list = [
        ['12/09/2024', '14:30', 'Jean', 'O', 'Ouverture normale', 'Succès'],
        ['12/09/2024', '18:00', 'Marie', 'F', 'Fermeture', 'Succès'],
        ['13/09/2024', '08:15', 'Pierre', 'DD', 'Défaillance détectée', 'IC']
    ]
    headers_list = ['Date', 'Heure', 'Opérateur', 'O/F/DD', 'Opération', 'Mention']
    
    # 3. Liste simple
    data_simple = ['Pomme', 'Banane', 'Orange', 'Fraise']
    
    # Créer l'exporteur
    exporter = DataExporter(output_dir="mes_exports")
    
    # Export complet (tous formats)
    print("=== Export données dictionnaire ===")
    files1 = exporter.export_data(
        data_dict, 
        "rapport_employes", 
        formats=['pdf', 'xlsx', 'docx'],
        title="Liste des Employés"
    )
    
    print("\n=== Export données liste avec headers ===")
    files2 = exporter.export_data(
        data_list,
        "BCC",
        formats=['pdf', 'xlsx', "docx"],  # Corrigé 'docs' en 'docx'
        title="BCC controle",
        headers=headers_list,
    )
    
    print("\n=== Export liste simple ===")
    files3 = exporter.export_data(
        data_simple,
        "liste_fruits",
        formats=['pdf'],
        title="Liste des Fruits"
    )
    
    print(f"\nFichiers créés: {files1}")
    print(f"Fichiers créés: {files2}")
    print(f"Fichiers créés: {files3}")
